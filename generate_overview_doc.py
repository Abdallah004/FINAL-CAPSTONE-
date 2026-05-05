"""Generate a non-technical Word document overview of the WarehouseMS project."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

ACCENT = RGBColor(0xE6, 0x39, 0x46)  # WarehouseMS red

doc = Document()

# Document defaults
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)


def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = ACCENT
    return h


def add_para(text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p


def add_bullet(text):
    p = doc.add_paragraph(text, style='List Bullet')
    return p


def add_kv(label, value):
    p = doc.add_paragraph()
    run = p.add_run(label + ': ')
    run.bold = True
    p.add_run(value)
    return p


# === TITLE PAGE ===
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('WarehouseMS')
run.bold = True
run.font.size = Pt(36)
run.font.color.rgb = ACCENT

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('A Web-Based Warehouse and Sales Management System')
run.font.size = Pt(14)
run.italic = True

doc.add_paragraph()
desc = doc.add_paragraph()
desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = desc.add_run('Project Overview — Non-Technical Guide')
run.font.size = Pt(12)

doc.add_page_break()

# === SECTION 1: WHAT IS IT ===
add_heading('1. What Is WarehouseMS?', level=1)

add_para(
    'WarehouseMS is a web application designed to help small supermarkets and '
    'food suppliers run their day-to-day business. Instead of juggling '
    'spreadsheets, paper notebooks, and separate apps for stock and sales, '
    'everything lives in one place. You open a website, log in, and you can '
    'see exactly what is on the shelves, what was sold today, what is running '
    'low, and what is about to expire.'
)

add_para(
    'It was built around the realities of a Lebanese supermarket — the kinds '
    'of products people actually buy here (Akkawi cheese, Almaza beer, '
    'manakeesh, kibbeh ingredients, and so on). It includes a smart '
    'recommendation feature that knows traditional Lebanese pairings, so it '
    'can suggest items that go together at the till.'
)

# === SECTION 2: WHO USES IT ===
add_heading('2. Who Uses It?', level=1)

add_para('The system has two types of users:')
add_para('Admin', bold=True)
add_para(
    'The owner or manager. Admins can do everything — add new products, '
    'create staff accounts, view all sales across the whole shop, download '
    'monthly reports, manage stock levels, and adjust inventory when there '
    'are corrections to make (for example, after a stock count).'
)
add_para('Staff', bold=True)
add_para(
    'The cashiers and floor employees. Staff can ring up sales, view the '
    'stock, generate invoices for customers, and see the alerts. They cannot '
    'see other staff members\' sales, change product prices, or access '
    'anything sensitive — those are admin-only powers.'
)

# === SECTION 3: WHAT IT DOES ===
add_heading('3. What Can It Do?', level=1)

add_heading('3.1 Inventory and Products', level=2)
add_para(
    'The system keeps a live, accurate picture of every product on the '
    'shelves. For each product you can see:'
)
add_bullet('The product name and description (e.g. "Akkawi Cheese 500 g — Mild white brine cheese")')
add_bullet('Its category (Dairy, Bakery, Produce, or Beverages)')
add_bullet('Current price')
add_bullet('How much is in stock right now')
add_bullet('When the next batch expires')
add_bullet('Whether stock is OK, running low, or completely out')

add_para(
    'There are 303 Lebanese-supermarket products pre-loaded — from Laban '
    'Ayran and Halloumi to Almaza beer, Château Musar wine, fresh figs, '
    'maamoul, and Najjar coffee. New products can be added at any time.'
)

add_heading('3.2 Sales (the cash register)', level=2)
add_para(
    'The "New Sale" page works like a digital cash register. The cashier '
    'searches for products, adds them to a cart, sets the quantity, and '
    'finalises the sale. Three things then happen automatically and at the '
    'same time:'
)
add_bullet('The customer\'s total is calculated.')
add_bullet('The sold quantity is deducted from inventory immediately.')
add_bullet('A printable PDF invoice is generated and can be downloaded.')

add_para(
    'Because all of that happens together, it is impossible for two cashiers '
    'to accidentally sell the same last item twice — the database refuses '
    'to oversell.'
)

add_heading('3.3 Smart Suggestions (the AI feature)', level=2)
add_para(
    'When a cashier adds items to the cart, a red "Smart Suggestions" panel '
    'appears with up to 5 recommended add-ons. The recommendations come from '
    'two places working together:'
)
add_para('Historical sales patterns', bold=True)
add_para(
    'The system looks at every past sale and asks "what other products were '
    'sold in the same basket as the items currently in this cart?" The more '
    'sales the shop processes, the smarter this gets — it learns your '
    'customers\' real habits.'
)
add_para('Lebanese cuisine pairings', bold=True)
add_para(
    'When sales history is sparse (or for a brand-new combination), the '
    'system uses around 30 hand-curated Lebanese food rules to fill in '
    'suggestions. Some examples:'
)
add_bullet('Hummus or foul → suggests pita, olive oil, pickles, mint')
add_bullet('Kibbeh or shawarma → suggests ayran, laban, pita, hot chili, tahini')
add_bullet('Manakeesh zaatar → suggests ayran, olives, cucumber, tea')
add_bullet('Maamoul or baklava → suggests Najjar coffee, cardamom coffee, sage tea')
add_bullet('Akkawi or halloumi cheese → suggests pita, tomato, mint, watermelon')
add_bullet('Lebanese beer (Almaza, 961) → suggests olives, pistachios, almonds, cheese')
add_bullet('Lebanese wine (Musar, Ksara, Kefraya) → suggests cheese, walnuts, figs, dates')

add_para(
    'A cashier can click any suggestion to add it straight to the cart. The '
    'goal is to gently boost basket size while making sense culturally — not '
    'to push random items.'
)

add_heading('3.4 Alerts', level=2)
add_para('The system watches stock and expiry dates around the clock and raises two kinds of alerts:')
add_bullet('Low-stock alerts — "X is running low" when the quantity drops to or below the threshold')
add_bullet('Expiry alerts — "X expires in N days" when items are due within 30 days')
add_para(
    'A bell icon in the header shows how many unread alerts there are. '
    'The system runs an automatic check every morning at 6 a.m. so the alerts '
    'are fresh when staff arrive.'
)

add_heading('3.5 Reports and Invoices', level=2)
add_para('Two types of PDF can be generated:')
add_bullet('Per-sale invoice — itemised receipt for any individual sale, ready to print or email')
add_bullet('Monthly sales report — overall KPIs, top-selling products, recent sales (admin only)')

add_heading('3.6 Activity Log', level=2)
add_para(
    'Every meaningful action — logins, sales, product changes, user '
    'creations, stock adjustments — is recorded in an audit log. Admins can '
    'see who did what and when, which is useful for reviews, training, and '
    'investigating any irregularity. Staff can see their own activity.'
)

add_heading('3.7 Dashboard', level=2)
add_para('The home page shows the day\'s key numbers at a glance:')
add_bullet('Total products in the catalog')
add_bullet('How many items are running low')
add_bullet('Today\'s revenue and the month\'s revenue')
add_bullet('Top-selling products and recent alerts')

# === SECTION 4: HOW IT IS BUILT ===
add_heading('4. How It Is Organized', level=1)

add_para('At a high level, WarehouseMS has three parts:')

add_para('The website (what you see)', bold=True)
add_para(
    'A clean, dark-themed dashboard that runs in any modern web browser. No '
    'app to install, no updates to manage. It works on laptops and desktops; '
    'tablets work but are not the primary target.'
)

add_para('The brain (the server)', bold=True)
add_para(
    'A program running in the cloud that handles every request from the '
    'website — checking passwords, calculating totals, deducting stock, '
    'creating PDFs, and so on. The cashier never sees this directly; it '
    'simply makes the website work.'
)

add_para('The memory (the database)', bold=True)
add_para(
    'A secure cloud database that stores every product, every sale, every '
    'user account, and every alert. Even if the website is closed or the '
    'computer is shut down, nothing is lost.'
)

# === SECTION 5: HOSTING ===
add_heading('5. Where It Lives', level=1)

add_para(
    'The system is hosted on professional cloud platforms so it does not '
    'depend on a computer in the shop:'
)
add_bullet('The website and brain run on Vercel — the same platform companies like TikTok and Twitch use.')
add_bullet('The database runs on Neon — a managed PostgreSQL service designed for reliability.')
add_para(
    'Both have automatic backups, encryption, and 24/7 monitoring built in. '
    'The shop never needs to install or maintain a server.'
)

# === SECTION 6: SECURITY ===
add_heading('6. Security and Access', level=1)
add_bullet('Passwords are never stored as plain text — they are scrambled (hashed) so even an attacker who sees the database cannot read them.')
add_bullet('Each session uses temporary, expiring tokens. If a phone or computer is lost, the session expires automatically.')
add_bullet('Admin features are completely off-limits to staff accounts — the server enforces this on every single request.')
add_bullet('All traffic between the browser and the server is encrypted (HTTPS).')
add_bullet('A rate limiter blocks anyone trying to guess passwords by hammering the login page.')

# === SECTION 7: WHY THIS MATTERS ===
add_heading('7. Why This Matters for the Business', level=1)

add_para('Compared to running on spreadsheets and paper, WarehouseMS:')
add_bullet('Eliminates double-entry — one sale updates the stock, the revenue, and the audit log in one click.')
add_bullet('Prevents overselling — the system refuses to sell items that are no longer in stock.')
add_bullet('Reduces waste — expiry alerts give the team time to discount or rotate stock before items go off.')
add_bullet('Increases basket size — Smart Suggestions surface items the customer is likely to want.')
add_bullet('Makes monthly reporting effortless — a single PDF replaces a day of manual calculations.')
add_bullet('Provides accountability — the activity log records who did what, which protects honest staff and deters mistakes.')

# === SECTION 8: GLOSSARY ===
add_heading('8. Glossary', level=1)
add_kv('Inventory', 'The list of every product the shop sells, with how many of each are currently on the shelf.')
add_kv('Threshold', 'The quantity at which a product is considered "running low" and triggers an alert. Different products can have different thresholds.')
add_kv('Sale', 'A complete transaction with a customer — one or more products, a total, and a timestamp.')
add_kv('Alert', 'An automatic warning generated when a product is running low or about to expire.')
add_kv('Admin', 'A user with full access to the system, including creating staff accounts and managing prices.')
add_kv('Staff', 'A user with limited access who can ring up sales but cannot change products, prices, or other users.')
add_kv('Cart', 'The list of items the cashier has added but not yet finalised as a sale.')
add_kv('Smart Suggestions', 'The AI feature that recommends related products to add to the cart.')
add_kv('Dashboard', 'The home page that shows the day\'s key numbers at a glance.')
add_kv('PDF Invoice', 'A printable receipt for a sale that can be saved or emailed to the customer.')

# === FOOTER ===
doc.add_paragraph()
foot = doc.add_paragraph()
foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = foot.add_run('WarehouseMS — One place for inventory and sales.')
run.italic = True
run.font.size = Pt(10)


output = r'C:\Users\Lenovo\OneDrive\Desktop\AI SYSTEM\WarehouseMS-Overview.docx'
doc.save(output)
print(f'Created: {output}')
